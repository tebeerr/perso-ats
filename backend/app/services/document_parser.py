import io
import re
from dataclasses import dataclass
from pathlib import Path

import fitz
from docx import Document

MAX_FILE_SIZE = 10 * 1024 * 1024
DOCX_MAGIC = b"PK\x03\x04"
EXPECTED_MIME_TYPES = {
    ".pdf": {"application/pdf"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
}


class DocumentError(ValueError):
    pass


@dataclass
class ParsedDocument:
    filename: str
    file_type: str
    text: str


def _clean_filename(filename: str) -> str:
    return Path(filename).name.replace("\x00", "")[:255]


def parse_document(filename: str, content_type: str | None, raw: bytes) -> ParsedDocument:
    clean_name = _clean_filename(filename)
    extension = Path(clean_name).suffix.lower()
    if not clean_name or extension not in {".pdf", ".docx"}:
        raise DocumentError("Only PDF and DOCX files are supported.")
    if content_type and content_type not in EXPECTED_MIME_TYPES[extension] | {"application/octet-stream"}:
        raise DocumentError("The file MIME type does not match its extension.")
    if len(raw) == 0:
        raise DocumentError("The uploaded document is empty.")
    if len(raw) > MAX_FILE_SIZE:
        raise DocumentError("The document exceeds the 10 MB limit.")
    if extension == ".pdf":
        if not raw.startswith(b"%PDF"):
            raise DocumentError("The file does not contain a valid PDF signature.")
        text = _parse_pdf(raw)
        file_type = "pdf"
    else:
        if not raw.startswith(DOCX_MAGIC):
            raise DocumentError("The file does not contain a valid DOCX signature.")
        text = _parse_docx(raw)
        file_type = "docx"
    if len(text.strip()) < 20:
        raise DocumentError("Unable to extract readable text. The PDF may be scanned or the document may be empty.")
    return ParsedDocument(clean_name, file_type, text)


def _parse_pdf(raw: bytes) -> str:
    try:
        document = fitz.open(stream=raw, filetype="pdf")
        pages = [page.get_text("text").strip() for page in document]
        document.close()
        return "\n\n".join(page for page in pages if page)
    except Exception as error:
        raise DocumentError("Unable to read this PDF document.") from error


def _parse_docx(raw: bytes) -> str:
    try:
        document = Document(io.BytesIO(raw))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = [cell.text.strip() for table in document.tables for row in table.rows for cell in row.cells if cell.text.strip()]
        return "\n".join(paragraphs + table_cells)
    except Exception as error:
        raise DocumentError("Unable to read this DOCX document.") from error


def detect_possible_columns(text: str) -> bool:
    return bool(re.search(r"\S\s{6,}\S", text))
